"""
Universal document extraction supporting:
- PDF (text-based and scanned/image-based)
- DOCX/DOC (Word documents)
- Images (JPG, PNG, etc.)

Uses pdfplumber for text PDFs (fast), PyMuPDF + Tesseract OCR for scanned PDFs,
and python-docx for Word documents.
"""

import io
import os
import shutil
import subprocess
import shlex
import tempfile
from pathlib import Path
from typing import Optional

import pdfplumber
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
from docx import Document

# Configure Tesseract path - try multiple common installation locations
def _find_tesseract():
    """Find Tesseract executable in common Windows installation paths"""
    common_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Users\heman\AppData\Local\Tesseract-OCR\tesseract.exe',
    ]
    
    # Try common installation locations first (most likely on Windows)
    for path in common_paths:
        if os.path.exists(path):
            print(f"Found Tesseract at: {path}")
            return path
    
    # Try to find in PATH as fallback
    tesseract_path = shutil.which('tesseract')
    if tesseract_path:
        print(f"Found Tesseract in PATH at: {tesseract_path}")
        return tesseract_path
    
    return None

# Set Tesseract path if found
_tesseract_path = _find_tesseract()
if _tesseract_path:
    # Configure pytesseract to use the found path
    pytesseract.pytesseract.pytesseract_cmd = _tesseract_path
    print(f"Tesseract configured at: {_tesseract_path}")
    
    # Test if tesseract is accessible
    try:
        result = subprocess.run(
            [_tesseract_path, '--version'],
            capture_output=True,
            timeout=5,
            text=True,
            check=False
        )
        if result.returncode == 0:
            print(f"✓ Tesseract is accessible and working")
        else:
            print(f"⚠ Tesseract found but may have issues: {result.stderr}")
    except Exception as e:
        print(f"⚠ Could not verify Tesseract: {e}")
else:
    print("Warning: Tesseract not found. Install from: https://github.com/UB-Mannheim/tesseract/wiki")


def _ocr_image_via_tesseract(image: Image.Image) -> str:
    """
    Extract text from image using Tesseract subprocess directly.
    This bypasses pytesseract's path issues on Windows.
    """
    if not _tesseract_path:
        raise ValueError("Tesseract not available for OCR")
    
    try:
        # Save image to temporary file
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
            image.save(tmp_img.name, 'PNG')
            tmp_img_path = tmp_img.name
        
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp_out:
            tmp_out_path = tmp_out.name
        
        try:
            # Call tesseract directly: tesseract input.png output
            # (tesseract adds .txt extension automatically)
            cmd = [_tesseract_path, tmp_img_path, tmp_out_path[:-4]]  # Remove .txt for output base
            result = subprocess.run(
                cmd,
                capture_output=True,
                timeout=30,
                text=True,
                check=False
            )
            
            if result.returncode != 0:
                error_msg = result.stderr or result.stdout or "Unknown error"
                raise ValueError(f"Tesseract execution failed: {error_msg}")
            
            # Read the output text file
            output_file = tmp_out_path[:-4] + '.txt'
            if os.path.exists(output_file):
                with open(output_file, 'r', encoding='utf-8') as f:
                    text = f.read()
                os.remove(output_file)
                return text.strip()
            else:
                raise ValueError("Tesseract did not produce output file")
        finally:
            # Clean up temp files
            if os.path.exists(tmp_img_path):
                os.remove(tmp_img_path)
            if os.path.exists(tmp_out_path):
                try:
                    os.remove(tmp_out_path)
                except:
                    pass
    except subprocess.TimeoutExpired:
        raise ValueError("Tesseract OCR timed out (processing took >30 seconds)")
    except Exception as e:
        if "Tesseract execution failed" in str(e):
            raise e
        raise ValueError(f"OCR processing error: {e}")


def extract_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF.
    First tries pdfplumber (fast for text PDFs).
    Falls back to OCR if extraction is minimal (scanned PDFs).
    """
    try:
        # Try fast text extraction first
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
            
            # If we got reasonable amount of text, return it
            if len(text) > 100:
                return text
    except Exception as e:
        print(f"Warning: pdfplumber extraction failed: {e}")
    
    # Fallback to OCR for scanned PDFs using PyMuPDF
    if not _tesseract_path:
        raise ValueError(
            "Cannot extract scanned PDF - Tesseract OCR not installed.\n"
            "Install from: https://github.com/UB-Mannheim/tesseract/wiki"
        )
    
    print("Attempting OCR extraction (scanned PDF detected)...")
    try:
        # Open PDF with PyMuPDF
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        if pdf_doc.page_count == 0:
            raise ValueError("PDF contains no pages")
        
        ocr_text = ""
        for page_num in range(pdf_doc.page_count):
            print(f"Processing page {page_num + 1}/{pdf_doc.page_count} with OCR...")
            
            # Render page to image (2x zoom for better OCR)
            page = pdf_doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            
            # Convert pixmap to PIL Image
            img_data = pix.tobytes("ppm")
            img = Image.open(io.BytesIO(img_data))
            
            # Extract text using Tesseract via subprocess
            try:
                page_text = _ocr_image_via_tesseract(img)
                ocr_text += page_text + "\n"
            except ValueError as e:
                raise ValueError(f"OCR failed on page {page_num + 1}: {str(e)}")
        
        pdf_doc.close()
        
        result = ocr_text.strip()
        if not result or len(result) < 10:
            raise ValueError("OCR extraction produced minimal text - PDF may be empty or corrupted")
        return result
    except ValueError as e:
        raise e
    except Exception as e:
        raise ValueError(f"PDF OCR extraction failed: {e}")


def extract_from_image(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from image files (JPG, PNG, etc.) using OCR.
    """
    if not _tesseract_path:
        raise ValueError(
            "Tesseract OCR not installed. Install from: https://github.com/UB-Mannheim/tesseract/wiki"
        )
    
    try:
        image = Image.open(io.BytesIO(file_bytes))
        text = _ocr_image_via_tesseract(image)
        if not text or len(text.strip()) < 5:
            raise ValueError("Image appears to be empty or contains no text")
        return text.strip()
    except ValueError as e:
        raise e
    except Exception as e:
        raise ValueError(f"Image OCR extraction failed: {e}")


def extract_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from Word documents (.docx, .doc).
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        
        if not text or len(text.strip()) < 5:
            raise ValueError("Word document appears to be empty")
        return text.strip()
    except ValueError as e:
        raise e
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}. File may be corrupted or not a valid Word document.")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Universal text extraction for multiple document formats.
    
    Supports:
    - PDF files (text and scanned)
    - Word documents (.docx, .doc)
    - Images (.jpg, .jpeg, .png, .tiff, .bmp)
    
    Args:
        file_bytes: Binary content of the file
        filename: Original filename (used to determine format)
    
    Returns:
        Extracted text from the document
    
    Raises:
        ValueError: If extraction fails or format is unsupported
    """
    filename_lower = filename.lower()
    
    # Determine file type by extension
    if filename_lower.endswith('.pdf'):
        text = extract_from_pdf(file_bytes)
    elif filename_lower.endswith(('.docx', '.doc')):
        text = extract_from_docx(file_bytes)
    elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif')):
        text = extract_from_image(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file format: {filename}. Supported: PDF, DOCX, DOC, JPG, PNG, TIFF, BMP, GIF")
    
    if not text or len(text.strip()) < 10:
        raise ValueError("Could not extract meaningful text from document. File may be corrupted or empty.")
    
    return text